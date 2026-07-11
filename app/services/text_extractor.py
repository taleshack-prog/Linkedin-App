"""Extração de texto de arquivos de referência das pautas.

Suporta: PDF (pypdf), DOCX (python-docx), TXT/MD/CSV (decodificação direta).
O texto extraído é truncado e injetado no prompt de geração — o binário
original não é armazenado.
"""
import io

MAX_SOURCE_CHARS = 60_000  # teto do texto guardado/injetado no prompt

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv"}


class ExtractionError(Exception):
    pass


def _ext(filename: str) -> str:
    name = (filename or "").lower()
    dot = name.rfind(".")
    return name[dot:] if dot != -1 else ""


def extract_text(filename: str, data: bytes) -> str:
    """Extrai texto do arquivo. Levanta ExtractionError se não suportado/vazio."""
    ext = _ext(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        raise ExtractionError(
            f"Formato '{ext or 'desconhecido'}' não suportado — use PDF, DOCX, TXT, MD ou CSV"
        )

    if ext == ".pdf":
        from pypdf import PdfReader

        try:
            reader = PdfReader(io.BytesIO(data))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception as exc:  # noqa: BLE001
            raise ExtractionError(f"Falha ao ler o PDF: {exc}")
    elif ext == ".docx":
        from docx import Document

        try:
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(parts)
        except Exception as exc:  # noqa: BLE001
            raise ExtractionError(f"Falha ao ler o DOCX: {exc}")
    else:  # txt / md / csv
        text = data.decode("utf-8", errors="ignore")

    text = text.strip()
    if not text:
        raise ExtractionError("Não foi possível extrair texto do arquivo (está vazio ou é digitalizado sem OCR)")
    return text[:MAX_SOURCE_CHARS]
