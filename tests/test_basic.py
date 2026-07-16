"""Testes unitários das partes puras (sem rede, sem banco)."""
import os
from datetime import datetime, timedelta, timezone

# Env mínimo ANTES de importar app.* (Settings valida no primeiro get_settings())
os.environ.setdefault("SECRET_KEY", "x" * 43 + "=")  # placeholder; Fernet só é usado lazy
os.environ.setdefault("DATABASE_URL", "sqlite://")
os.environ.setdefault("LINKEDIN_CLIENT_ID", "test")
os.environ.setdefault("LINKEDIN_CLIENT_SECRET", "test")
os.environ.setdefault("ANTHROPIC_API_KEY", "test")

import pytest
from pydantic import ValidationError

from app.schemas import PostApprove, PostUpdate
from app.services.linkedin_client import escape_commentary


class TestEscapeCommentary:
    def test_escapa_reservados_do_little_text(self):
        assert escape_commentary("a*b_c(d)") == r"a\*b\_c\(d\)"

    def test_backslash_escapado_primeiro(self):
        # Se '\' não fosse o primeiro, escaparíamos os escapes gerados.
        assert escape_commentary("a\\b") == "a\\\\b"

    def test_texto_limpo_intacto(self):
        assert escape_commentary("Post normal, sem reservados.") == "Post normal, sem reservados."


class TestPostApprove:
    def test_rejeita_publish_at_no_passado(self):
        with pytest.raises(ValidationError):
            PostApprove(publish_at=datetime.now(timezone.utc) - timedelta(hours=1))

    def test_naive_vira_utc(self):
        futuro_naive = datetime.utcnow() + timedelta(days=1)  # noqa: DTZ003
        approved = PostApprove(publish_at=futuro_naive)
        assert approved.publish_at.tzinfo is not None

    def test_aceita_futuro_aware(self):
        alvo = datetime.now(timezone.utc) + timedelta(hours=2)
        assert PostApprove(publish_at=alvo).publish_at == alvo


class TestPostUpdate:
    def test_rejeita_commentary_acima_de_3000(self):
        with pytest.raises(ValidationError):
            PostUpdate(commentary="x" * 3001)

    def test_aceita_commentary_no_limite(self):
        assert len(PostUpdate(commentary="x" * 3000).commentary) == 3000


class TestFernetLazy:
    def test_import_nao_exige_fernet_valido(self):
        # security foi importado indiretamente sem instanciar Fernet — se
        # chegou aqui, o lazy init funciona.
        import app.security  # noqa: F401


class TestPostPayload:
    def test_payload_sem_imagem_nao_tem_content(self):
        from app.services.linkedin_client import build_post_payload
        p = build_post_payload("urn:li:person:X", "texto")
        assert "content" not in p and p["author"] == "urn:li:person:X"

    def test_payload_com_imagem_referencia_urn(self):
        from app.services.linkedin_client import build_post_payload
        p = build_post_payload("urn:li:person:X", "texto", image_urn="urn:li:image:ABC")
        assert p["content"] == {"media": {"id": "urn:li:image:ABC"}}

    def test_commentary_escapado_no_payload(self):
        from app.services.linkedin_client import build_post_payload
        p = build_post_payload("urn:li:person:X", "a(b)")
        assert p["commentary"] == r"a\(b\)"


class TestImageGenerator:
    def test_prompt_inclui_tema_e_proibe_texto(self):
        from app.services.image_generator import build_image_prompt
        p = build_image_prompt("Tendências de Web3 no Brasil")
        assert "Tendências de Web3 no Brasil" in p and "NÃO incluir nenhum texto" in p

    def test_prompt_trunca_e_anexa_instrucoes(self):
        from app.services.image_generator import build_image_prompt
        p = build_image_prompt("x" * 5000, instructions="tons de azul")
        assert "x" * 700 in p and "x" * 701 not in p and "tons de azul" in p

    def test_parse_extrai_base64_e_mime(self):
        import base64
        from app.services.image_generator import parse_image_response
        raw = b"\x89PNG-fake"
        data = {"candidates": [{"content": {"parts": [
            {"text": "aqui está"},
            {"inlineData": {"mimeType": "image/png", "data": base64.b64encode(raw).decode()}},
        ]}}]}
        img, mime = parse_image_response(data)
        assert img == raw and mime == "image/png"

    def test_parse_sem_imagem_levanta_erro(self):
        import pytest as _pytest
        from app.services.image_generator import ImageGenError, parse_image_response
        with _pytest.raises(ImageGenError):
            parse_image_response({"candidates": [{"content": {"parts": [{"text": "só texto"}]}}]})


class TestProfileContext:
    def test_perfil_vazio_gera_bloco_vazio(self):
        from app.services.content_generator import build_profile_context
        assert build_profile_context(None) == ""
        assert build_profile_context({"role": None, "goal": None}) == ""

    def test_contexto_traduz_atuacao_e_objetivo(self):
        from app.services.content_generator import build_profile_context
        ctx = build_profile_context({
            "entity_type": "autonomo", "role": "eng. de software", "company": "HTF",
            "industry": "tecnologia", "audience": "fundadores", "goal": "leads",
            "tone": "direto", "pillars": "IA; automação", "positioning": "jogo infinito",
        })
        assert "profissional autônomo(a)" in ctx
        assert "gerar leads/clientes" in ctx
        assert "jogo infinito" in ctx and "Contexto do autor" in ctx

    def test_campos_ausentes_sao_omitidos(self):
        from app.services.content_generator import build_profile_context
        ctx = build_profile_context({"goal": "autoridade"})
        assert "Objetivo" in ctx and "Tom de voz" not in ctx and "Empresa" not in ctx


class TestOpenAIProvider:
    def test_parse_openai_extrai_b64(self):
        import base64
        from app.services.image_generator import parse_openai_response
        raw = b"PNG-openai"
        img, mime = parse_openai_response({"data": [{"b64_json": base64.b64encode(raw).decode()}]})
        assert img == raw and mime == "image/png"

    def test_parse_openai_sem_dados_levanta_erro(self):
        import pytest as _pytest
        from app.services.image_generator import ImageGenError, parse_openai_response
        with _pytest.raises(ImageGenError):
            parse_openai_response({"data": []})
        with _pytest.raises(ImageGenError):
            parse_openai_response({"data": [{"b64_json": None}]})

    def test_provider_invalido_levanta_503(self):
        import pytest as _pytest
        from app.config import get_settings
        from app.services.image_generator import ImageGenError, generate_post_image
        get_settings().IMAGE_PROVIDER = "midjourney"
        try:
            with _pytest.raises(ImageGenError) as exc:
                generate_post_image("post qualquer")
            assert exc.value.status == 503
        finally:
            get_settings().IMAGE_PROVIDER = "gemini"


class TestTextExtractor:
    def test_txt_extrai_e_trunca(self):
        from app.services.text_extractor import MAX_SOURCE_CHARS, extract_text
        assert extract_text("nota.txt", "relatório Q3 da HTF".encode()) == "relatório Q3 da HTF"
        assert len(extract_text("big.md", b"a" * 100_000)) == MAX_SOURCE_CHARS

    def test_docx_extrai_paragrafos_e_tabelas(self):
        import io
        from docx import Document
        from app.services.text_extractor import extract_text
        doc = Document()
        doc.add_paragraph("Resultados do trimestre")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Receita"
        table.rows[0].cells[1].text = "R$ 100k"
        buf = io.BytesIO(); doc.save(buf)
        text = extract_text("relatorio.docx", buf.getvalue())
        assert "Resultados do trimestre" in text and "Receita | R$ 100k" in text

    def test_formato_nao_suportado_e_vazio(self):
        import pytest as _pytest
        from app.services.text_extractor import ExtractionError, extract_text
        with _pytest.raises(ExtractionError):
            extract_text("virus.exe", b"MZ")
        with _pytest.raises(ExtractionError):
            extract_text("vazio.txt", b"   ")


class TestSourceBlock:
    def test_bloco_orienta_basear_no_material(self):
        from app.services.content_generator import build_source_block
        block = build_source_block("conteúdo do relatório")
        assert "PRINCIPALMENTE" in block and "conteúdo do relatório" in block
        assert build_source_block(None) == "" and build_source_block("") == ""


class TestRobustJsonParsing:
    def test_json_com_quebra_literal_dentro_da_string(self):
        # Reproduz o bug real: "Expecting ',' delimiter" com \n literal no commentary
        from app.services.content_generator import parse_json_lenient
        raw = '{"posts": [{"commentary": "linha 1\nlinha 2\n\n\\"citação\\" ok", "hashtags": []}]}'
        data = parse_json_lenient("bla bla " + raw + " tchau")
        assert data["posts"][0]["commentary"] == 'linha 1\nlinha 2\n\n"citação" ok'

    def test_tool_use_tem_prioridade_sobre_texto(self):
        from types import SimpleNamespace as NS
        from app.services.content_generator import extract_posts_payload
        msg = NS(content=[
            NS(type="text", text="pesquisando..."),
            NS(type="tool_use", name="emit_posts",
               input={"posts": [{"commentary": "post via tool", "hashtags": ["ia"]}]}),
        ])
        assert extract_posts_payload(msg)["posts"][0]["commentary"] == "post via tool"

    def test_fallback_para_texto_quando_nao_ha_tool(self):
        from types import SimpleNamespace as NS
        from app.services.content_generator import extract_posts_payload
        msg = NS(content=[NS(type="text", text='{"posts": [{"commentary": "via texto"}]}')])
        assert extract_posts_payload(msg)["posts"][0]["commentary"] == "via texto"

    def test_schema_da_tool_exige_commentary(self):
        from app.services.content_generator import POSTS_TOOL
        item = POSTS_TOOL["input_schema"]["properties"]["posts"]["items"]
        assert "commentary" in item["required"] and POSTS_TOOL["name"] == "emit_posts"


class TestPromptNaoSequestraTema:
    def test_system_prompt_prioriza_tema_da_pauta(self):
        from app.services.content_generator import SYSTEM_PROMPT
        assert "O TEMA DA PAUTA sempre manda" in SYSTEM_PROMPT
        assert "NUNCA desvie o assunto" in SYSTEM_PROMPT

    def test_cabecalho_do_contexto_reflete_a_regra(self):
        from app.services.content_generator import build_profile_context
        ctx = build_profile_context({"pillars": "Web3"})
        assert "o tema da pauta manda" in ctx


class TestAuthJWT:
    def test_bcrypt_hash_e_verificacao(self):
        from app.security import hash_password, verify_password
        h = hash_password("minhasenha123")
        assert h != "minhasenha123" and verify_password("minhasenha123", h)
        assert not verify_password("errada", h)

    def test_jwt_roundtrip_e_token_invalido(self):
        import pytest as _pytest
        from fastapi import HTTPException
        from app.security import create_token, decode_token
        assert decode_token(create_token("abc-123")) == "abc-123"
        with _pytest.raises(HTTPException):
            decode_token("token.invalido.aqui")


class TestPlansAndReferrals:
    def test_escada_de_meses(self):
        from app.services.plans import months_earned
        assert [months_earned(n) for n in (0, 2, 3, 9, 10, 15, 16, 50)] == [0, 0, 1, 1, 6, 6, 12, 12]

    def test_gating_por_plano(self):
        from app.services.plans import PLANS
        assert PLANS["starter"].ai_images is False and PLANS["starter"].doc_upload is False
        assert PLANS["pro"].ai_images is True and PLANS["pro"].doc_upload is True
        assert PLANS["free"].linkedin_accounts == 1 and PLANS["agency"].linkedin_accounts == 10
        assert [PLANS[k].price_cents for k in ("starter", "pro", "agency")] == [2000, 4570, 10000]
        # formatação de texto: Pro+ apenas
        assert PLANS["starter"].text_formatting is False
        assert PLANS["pro"].text_formatting is True and PLANS["agency"].text_formatting is True

    def test_plano_expirado_vira_free(self):
        from datetime import datetime, timezone
        from app.services.plans import plan_of
        class U:
            plan = "pro"
            plan_until = datetime(2020, 1, 1, tzinfo=timezone.utc)
        assert plan_of(U()).key == "free"

    def test_plano_sem_validade_vale(self):
        from app.services.plans import plan_of
        class U:
            plan = "starter"
            plan_until = None
        assert plan_of(U()).key == "starter"


class TestReferredBonus:
    def test_bonus_do_indicado_definido(self):
        from app.services.plans import REFERRED_BONUS_DAYS
        assert REFERRED_BONUS_DAYS == 15


def _so(d):
    """Constrói um StripeObject igual ao que a lib entrega no webhook.
    Testar com dict puro NÃO reproduz produção: StripeObject não tem .get()."""
    from stripe._stripe_object import StripeObject
    return StripeObject.construct_from(d, "sk_test")


class TestStripeWebhookCompat:
    """Dois riscos reais cobertos aqui:
    1) API Basil (2025-03-31) moveu current_period_end para os items;
    2) StripeObject (lib v15+) não herda dict e não tem .get()."""

    def test_acessor_funciona_com_stripeobject_e_dict(self):
        from app.routers.billing import _g
        so = _so({"customer": "cus_1", "metadata": {"plan": "pro"}})
        assert _g(so, "customer") == "cus_1"           # StripeObject (produção)
        assert _g({"customer": "cus_2"}, "customer") == "cus_2"  # dict
        assert _g(so, "inexistente", "PADRAO") == "PADRAO"
        assert _g(None, "x", "PADRAO") == "PADRAO"
        assert _g(_g(so, "metadata"), "plan") == "pro"  # aninhado

    def test_period_end_com_stripeobject_basil(self):
        from app.routers.billing import _period_end
        d = _period_end(_so({"items": {"data": [{"current_period_end": 1800000000}]}}))
        assert d is not None and d.year == 2027

    def test_period_end_formato_novo_basil(self):
        from app.routers.billing import _period_end
        sub = {"items": {"data": [{"current_period_end": 1800000000}]}}
        d = _period_end(sub)
        assert d is not None and d.year == 2027

    def test_period_end_formato_antigo(self):
        from app.routers.billing import _period_end
        assert _period_end({"current_period_end": 1800000000}) is not None

    def test_period_end_ausente_nao_quebra(self):
        from app.routers.billing import _period_end
        assert _period_end({}) is None
        assert _period_end({"items": {"data": []}}) is None

    def test_subscription_id_da_invoice_com_stripeobject(self):
        from app.routers.billing import _subscription_id_from_invoice
        novo = _so({"parent": {"subscription_details": {"subscription": "sub_SO"}}})
        assert _subscription_id_from_invoice(novo) == "sub_SO"
        assert _subscription_id_from_invoice(_so({})) is None

    def test_subscription_id_da_invoice_novo_e_antigo(self):
        from app.routers.billing import _subscription_id_from_invoice
        # formato antigo
        assert _subscription_id_from_invoice({"subscription": "sub_123"}) == "sub_123"
        # formato novo (Basil): parent.subscription_details.subscription
        novo = {"parent": {"subscription_details": {"subscription": "sub_456"}}}
        assert _subscription_id_from_invoice(novo) == "sub_456"
        # invoice avulsa
        assert _subscription_id_from_invoice({}) is None


class TestLGPD:
    def test_relacionamento_delega_cascade_ao_banco(self):
        """Sem passive_deletes o ORM tenta user_id=NULL (NOT NULL) e a exclusão
        de conta quebra — violando o art. 18, VI. Bug real, pego em teste."""
        from app.models import User
        rel = User.__mapper__.relationships["linkedin_accounts"]
        assert rel.passive_deletes is True
        assert "delete-orphan" in rel.cascade

    def test_confirmacao_de_exclusao_e_case_insensitive(self):
        # 'excluir', 'EXCLUIR' e ' Excluir ' devem valer; outros textos não
        for ok in ("EXCLUIR", "excluir", "  Excluir  "):
            assert ok.strip().upper() == "EXCLUIR"
        for nao in ("sim", "delete", "", "exclui"):
            assert nao.strip().upper() != "EXCLUIR"


class TestViradaTestLive:
    """Customer criado em modo teste não existe em live. Sem tratar, o primeiro
    checkout em produção quebra com 'No such customer' para quem já testou."""

    def test_recria_customer_invalido_e_reutiliza_valido(self):
        from app.routers.billing import _ensure_customer

        class FakeDB:
            def commit(self): pass

        class FakeUser:
            email = "x@t.com"
            id = "uid"
            stripe_customer_id = "cus_TEST_velho"

        class S:
            class Customer:
                @staticmethod
                def retrieve(cid):
                    if cid.startswith("cus_TEST"):
                        raise Exception("No such customer")
                    class C: deleted = False
                    return C()
                @staticmethod
                def create(email, metadata):
                    class C: id = "cus_NOVO"
                    return C()

        u = FakeUser()
        assert _ensure_customer(S, FakeDB(), u) == "cus_NOVO"   # inválido -> recria
        assert u.stripe_customer_id == "cus_NOVO"               # e persiste
        u.stripe_customer_id = "cus_LIVE_ok"
        assert _ensure_customer(S, FakeDB(), u) == "cus_LIVE_ok"  # válido -> reutiliza


class TestPlanoAnual:
    def test_anual_e_10x_o_mensal_com_2_meses_gratis(self):
        from app.services.plans import PLANS, annual_savings_cents
        for k in ("starter", "pro", "agency"):
            p = PLANS[k]
            assert p.price_cents_annual == p.price_cents * 10, k
            assert annual_savings_cents(p) == p.price_cents * 2, k  # economia = 2 meses

    def test_valores_anuais(self):
        from app.services.plans import PLANS
        assert PLANS["starter"].price_cents_annual == 20000     # R$ 200,00
        assert PLANS["pro"].price_cents_annual == 45700         # R$ 457,00
        assert PLANS["agency"].price_cents_annual == 100000     # R$ 1.000,00

    def test_price_id_escolhe_pelo_ciclo(self):
        from app.config import get_settings
        from app.routers.billing import _price_id
        s = get_settings()
        s.STRIPE_PRICE_PRO = "price_mensal"
        s.STRIPE_PRICE_PRO_ANNUAL = "price_anual"
        try:
            assert _price_id("pro", "monthly") == "price_mensal"
            assert _price_id("pro", "annual") == "price_anual"
            assert _price_id("pro") == "price_mensal"        # default
            assert _price_id("inexistente", "annual") is None
        finally:
            s.STRIPE_PRICE_PRO = ""
            s.STRIPE_PRICE_PRO_ANNUAL = ""


class TestLandingPublica:
    def test_plans_e_publico_para_a_landing(self):
        """A landing consulta /billing/plans SEM login — se exigisse auth,
        o visitante não veria preço nenhum."""
        import inspect
        from app.routers import billing
        src = inspect.getsource(billing.list_plans)
        assert "get_current_user" not in src
